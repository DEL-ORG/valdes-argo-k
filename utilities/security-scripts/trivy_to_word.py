#!/usr/bin/env python3
"""
Trivy Scan Automation
---------------------
Runs Trivy scan, generates Word report, posts summary to Mattermost,
and uploads report to S3 when running on a tagged branch.

Usage:
  python3 utilities/security-scripts/trivy_to_word.py \
    --image myrepo/myimage:latest \
    --branch main
"""

import argparse, os, sys, json, subprocess, logging, requests, boto3
from datetime import datetime
from pathlib import Path
from docx import Document

# === CONFIG ===
S3_BUCKET = os.getenv("S3_BUCKET", "security-trivy-reports")
ARTIFACTS_DIR = Path("artifacts")
MATTERMOST_WEBHOOK = os.getenv("MATTERMOST_WEBHOOK_URL")
AWS_REGION = os.getenv("AWS_REGION", "us-east-1")
TRIVY = os.getenv("TRIVY_BINARY", "trivy")

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

def run(cmd):
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode not in (0, 1):
        logging.error("Command failed: %s", res.stderr)
        sys.exit(1)
    return res.stdout

def update_trivy_db():
    logging.info("Updating Trivy DB...")
    # compatible with new and old Trivy versions
    try:
        subprocess.run([TRIVY, "image", "--download-db-only"], check=True)
    except subprocess.CalledProcessError:
        logging.warning("`--download-db-only` not supported, using `--download-db-only` fallback...")
        subprocess.run([TRIVY, "image", "--update"], check=False)


def run_trivy_json(image):
    logging.info(f"Scanning {image} ...")
    out = run([TRIVY, "image", "--format", "json", "--severity", "CRITICAL,HIGH", image])
    return json.loads(out)

def extract_vulns(data):
    vulns = []
    for r in data.get("Results", []):
        for v in r.get("Vulnerabilities", []) or []:
            vulns.append({
                "ID": v.get("VulnerabilityID", ""),
                "Pkg": v.get("PkgName", ""),
                "Severity": v.get("Severity", ""),
                "Installed": v.get("InstalledVersion", ""),
                "Fixed": v.get("FixedVersion", ""),
                "Desc": (v.get("Description") or "")[:400],
            })
    return vulns

def make_doc(image, branch, vulns, ts, outpath):
    doc = Document()
    doc.add_heading(f"Trivy Report - {image}", 1)
    doc.add_paragraph(f"Branch/Tag: {branch}")
    doc.add_paragraph(f"Scan Time: {ts}")
    doc.add_paragraph(f"Total: {len(vulns)} (CRITICAL: {sum(v['Severity']=='CRITICAL' for v in vulns)}, HIGH: {sum(v['Severity']=='HIGH' for v in vulns)})")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    hdr = ["ID", "Package", "Severity", "Installed", "Fixed", "Description"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h

    for v in vulns:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = v["ID"], v["Pkg"], v["Severity"]
        cells[3].text, cells[4].text, cells[5].text = v["Installed"], v["Fixed"], v["Desc"]

    outpath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(outpath)
    logging.info(f"Saved report: {outpath}")

def post_mattermost(image, branch, ts, critical, high, s3_uri=None):
    if not MATTERMOST_WEBHOOK:
        logging.warning("MATTERMOST_WEBHOOK_URL not set — skipping Mattermost")
        return
    text = (
        f"**Trivy Scan Summary**\n"
        f"Image: `{image}`\nBranch/Tag: `{branch}`\nScan: {ts}\n"
        f"CRITICAL: {critical}, HIGH: {high}\n"
    )
    if s3_uri:
        text += f"Report: {s3_uri}"
    requests.post(MATTERMOST_WEBHOOK, json={"text": text}, timeout=10)
    logging.info("Summary posted to Mattermost.")

def upload_s3(path, bucket, key):
    s3 = boto3.client("s3", region_name=AWS_REGION)
    s3.upload_file(str(path), bucket, key)
    return f"s3://{bucket}/{key}"

def is_tag():
    if os.getenv("CI_COMMIT_TAG"):
        return True, os.getenv("CI_COMMIT_TAG")
    ref = os.getenv("GITHUB_REF", "")
    if ref.startswith("refs/tags/"):
        return True, ref.split("/")[-1]
    return False, None

def main():
    p = argparse.ArgumentParser()
    p.add_argument("--image", required=True, help="Container image to scan")
    p.add_argument("--branch", required=False, help="Branch or tag name")
    args = p.parse_args()

    branch = args.branch or os.getenv("CI_COMMIT_REF_NAME") or "unknown"
    tagged, tagname = is_tag()
    if tagged:
        branch = tagname

    ts = datetime.utcnow().strftime("%Y-%m-%dT%H-%M-%SZ")

    update_trivy_db()
    data = run_trivy_json(args.image)
    vulns = extract_vulns(data)

    critical = sum(v["Severity"] == "CRITICAL" for v in vulns)
    high = sum(v["Severity"] == "HIGH" for v in vulns)

    safe_img = args.image.replace("/", "_").replace(":", "_")
    report = ARTIFACTS_DIR / f"{safe_img}_{branch}_{ts}.docx"
    make_doc(args.image, branch, vulns, ts, report)

    s3_uri = None
    if tagged and os.getenv("AWS_ACCESS_KEY_ID"):
        key = f"{branch}/{safe_img}/{report.name}"
        s3_uri = upload_s3(report, S3_BUCKET, key)

    post_mattermost(args.image, branch, ts, critical, high, s3_uri)

    if critical + high > 0:
        logging.error("Found HIGH/CRITICAL vulnerabilities.")
        sys.exit(1)

if __name__ == "__main__":
    main()
